# -*- coding: utf-8 -*-
"""COURSE MASTER 仕様書 DOCX 生成スクリプト"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_PATH = os.path.join(BASE_DIR, '仕様書', 'COURSE_MASTER_仕様書_20260528.docx')

JP_FONT = 'MS Gothic'
EN_FONT = 'Arial'
HDR_BG  = 'D5E8F0'
TBL_BORDER = 'CCCCCC'

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_table_borders(table, color_hex):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), color_hex)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def set_run_font(run, size=10.5):
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), JP_FONT)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, 14 - level * 1.5)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # ヘッダー行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        set_run_font(run, 9.5)
        set_cell_bg(cell, HDR_BG)

    # データ行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, 9.5)

    set_table_borders(table, TBL_BORDER)

    # 列幅設定
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

# --- ドキュメント作成 ---
doc = Document()

# A4設定
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.0)

# デフォルトフォント
doc.styles['Normal'].font.name = EN_FONT
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), JP_FONT)

# ===== 表紙タイトル =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COURSE MASTER 仕様書')
run.bold = True
run.font.size = Pt(20)
run.font.name = EN_FONT
rPr = run._r.get_or_add_rPr()
rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), JP_FONT); rPr.insert(0, rf)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('バージョン 7.3 ／ 作成日 2026-05-28')
run2.font.size = Pt(11)
run2.font.name = EN_FONT
rPr2 = run2._r.get_or_add_rPr()
rf2 = OxmlElement('w:rFonts'); rf2.set(qn('w:eastAsia'), JP_FONT); rPr2.insert(0, rf2)

doc.add_paragraph()

# ===== 1. 概要 =====
add_heading(doc, '1. 概要', 1)
add_table(doc,
    ['項目', '内容'],
    [
        ['システム名',   'COURSE MASTER'],
        ['バージョン',   'v7.3（エンジン: course_master_v73_engine.py）'],
        ['目的',         'JRA重賞・条件戦の着順予測、市場歪みの特定（エッジ値算出）'],
        ['予測形式',     '◎1頭・○3頭・▲5頭を出力。7軸スコア＋4チーム合議制'],
        ['作成日',       '2026-05-28'],
        ['最終更新日',   '2026-05-28（月次メンテ・CSV取込・DBメンテ）'],
        ['管理場所',     'C:\\Users\\shiny\\Dropbox\\shinya_wa\\coursemaster\\'],
        ['Webアプリ',    'prediction_app.py（Flask・localhost:5000）'],
    ],
    col_widths=[4.0, 12.0]
)

# ===== 2. 蓄積データ現況 =====
doc.add_paragraph()
add_heading(doc, '2. 蓄積データ現況（2026-05-28 月次メンテ後）', 1)
add_table(doc,
    ['テーブル / 項目', '件数 / 値', '備考'],
    [
        ['race_results',  '551,656件', '2015-01-04〜2026-05-24'],
        ['blood_map',     '51,933件',  '血統マッピング（父馬・母の父）'],
        ['DBサイズ',      '115.47 MB', 'VACUUM後（旧123.81 MB）'],
        ['対象期間',      '2015〜2026年5月', '12年分 約11年半'],
        ['NULL率最大カラム', 'sire_name 23.6%', '30%超なし（正常範囲）'],
    ],
    col_widths=[5.0, 4.5, 6.5]
)
doc.add_paragraph()
add_heading(doc, 'DBテーブル一覧（統計マスター）', 2)
add_table(doc,
    ['テーブル名', '件数', '内容'],
    [
        ['sire_profile',        '431',    '種牡馬プロファイル（距離帯別勝率）'],
        ['bms_profile',         '730',    'BMS（母の父）プロファイル'],
        ['jockey_stats',        '249',    '騎手統計'],
        ['trainer_stats',       '261',    '調教師統計'],
        ['jockey_trainer_combo','10,865', '騎手×調教師コンビ統計'],
        ['horse_career',        '17,431', '馬別キャリア統計'],
        ['horse_speed',         '15,964', '馬別スピード指数'],
        ['speed_base',          '381',    '距離別スピード基準'],
        ['course_win_rate',     '403',    'コース別勝率'],
        ['field_bias',          '66',     'コース・馬場バイアス'],
        ['blood_category',      '431',    '血統カテゴリ（速力系/マイラー系/スタミナ系）'],
    ],
    col_widths=[4.5, 2.0, 9.5]
)

# ===== 3. エンジン仕様 =====
doc.add_paragraph()
add_heading(doc, '3. エンジン仕様（v7.3）', 1)

add_heading(doc, '3.1 スコアリング7軸（採用中）', 2)
add_table(doc,
    ['軸', '名称', '計算ベース', '重み', '備考'],
    [
        ['CF',  'キャリア形成',     '血統CSV全成績（勝率＋走数）',    '×2.0', '走数ペナルティあり（1〜3走目×0.70）'],
        ['SI',  'スピードインデックス', '上がり3Fタイム',             '×2.0', ''],
        ['JT',  'ジョッキー',       '騎手別統計勝率',                 '×2.0', ''],
        ['SPD', 'スピード能力',     '走破時計のz-score',              '×2.0', ''],
        ['PD',  'ペースデザイン',   '通過順（1〜4F）位置取り',        '×1.0', '欠損時50点'],
        ['BL',  'ベース力',         '人気順ベース評価',               '×0.3', ''],
        ['MK',  'マーケット',       '人気順＋オッズ帯補正',           '×0.3', '低オッズ帯を減点（2倍以下×0.85）'],
    ],
    col_widths=[1.5, 3.5, 5.0, 1.5, 5.0]
)
add_para(doc, '最終スコア = 7軸重み付き平均 × 4チーム（I/O/U/S）の平均値')

add_heading(doc, '3.2 除外軸（コード保持・スコア算出対象外）', 2)
add_table(doc,
    ['軸', '除外理由'],
    [
        ['FR', '前走馬体重データ非対応 → 常に50点固定'],
        ['CL', '同一レース内全馬が同じクラスコード → 全馬同値で差異なし'],
        ['TR', '同一レース内全馬が同じコース条件 → 全馬同値'],
        ['BF', '重賞1,544Rバックテストで着順との相関係数+0.020（逆効果）→ 2026-04-03除外'],
        ['HP', '重賞1,544Rバックテストで着順との相関係数-0.030（ほぼ無効）→ 2026-04-03除外'],
    ],
    col_widths=[2.0, 14.0]
)

add_heading(doc, '3.3 4チーム合議制', 2)
add_para(doc, 'I型（スタンダード）・O型（過学習抑制）・U型（穴馬重視）・S型（速力重視）の4チームが独立予測し平均を取る。')

add_heading(doc, '3.4 勝率変換（softmax）', 2)
add_para(doc, '温度パラメータ T=5.0 のsoftmax変換。P(win) <= P(place)/3 の原則を適用。')

add_heading(doc, '3.5 エッジ計算式', 2)
add_para(doc, 'エッジ値 = エンジン推定勝率 − 市場確率（1 / オッズ × 0.80）')
add_para(doc, '標準閾値: +0.06以上 → 購入推奨 ／ 阪神稍重時: +0.07以上')

add_heading(doc, '3.6 MK軸オッズ帯補正（v7.3 rev2）', 2)
add_table(doc,
    ['オッズ帯', 'MK補正係数', '根拠'],
    [
        ['〜2倍',  '×0.85', '実勝率 < 市場期待 −0.038（バックテスト2024-2026 Grade 277R）'],
        ['2〜3倍', '×0.90', '実勝率 < 市場期待 −0.043'],
        ['3倍超',  '×1.00', '補正なし'],
    ],
    col_widths=[3.0, 3.5, 9.5]
)

add_heading(doc, '3.7 v7.0→v7.3 主要変更点', 2)
add_table(doc,
    ['Fix', '内容'],
    [
        ['Fix1', 'CF/BF/CL/JT軸: place_rate補正・固定加算を廃止 → win_rate*100のみ'],
        ['Fix2', 'CF軸: キャリア走数ペナルティ導入（1〜3走目×0.70 / 4〜5走目×0.85）'],
        ['Fix3', 'MK軸: オッズ帯補正rev2（過大評価のみ減点、加点廃止）'],
        ['Phase3 Step5', 'BF/HP除外・7軸化 → 重賞的中率 32.2% → 35.3%（+3.1%改善）'],
    ],
    col_widths=[3.5, 12.5]
)

# ===== 4. 採用済みベットルール =====
doc.add_paragraph()
add_heading(doc, '4. 採用済みベットルール一覧', 1)

add_heading(doc, '4.1 重賞戦略（grade_race_predictor.py）', 2)
add_table(doc,
    ['ルール名', '条件', 'ROI / 精度', '検証R数', '採用日'],
    [
        ['穴馬戦略',
         '重賞 ◎4番人気以上 × 距離2200m未満 → 単勝購入',
         '227.7%（単勝）/ 黒字9/12年75%',
         '150R（2015〜2026）',
         '2026-04-15'],
        ['エッジ閾値（標準）',
         'エッジ+0.06以上 → 購入推奨（2200m以上除く）',
         '247.8%（検証回収率）',
         '320R（2024〜2026検証）',
         '2026-04-17'],
        ['エッジ閾値（阪神稍重）',
         'エッジ+0.07以上（阪神稍重のみ）',
         '100.0%（n=5・収支均衡）',
         '5R',
         '2026-04-17'],
        ['長距離ベット保留',
         '距離2200m以上: 予測実施・ベット保留。30R超で再評価',
         '累計2R◎的中2/2（検証中）',
         '2R（目標30R）',
         '2026-04-25'],
        ['穴馬×エッジ複合',
         '穴馬戦略＋エッジ+0.06両方該当 → ★★強推奨',
         '304.6%（複合条件）',
         '—',
         '2026-04-16'],
    ],
    col_widths=[3.5, 5.5, 3.5, 3.0, 2.5]
)

add_heading(doc, '4.2 条件戦戦略（A〜D）', 2)
add_table(doc,
    ['条件', '会場・条件', '血統', '人気', '見送り馬場', 'ROI', '安定性', '採用日'],
    [
        ['A（中京ダ）', '中京ダート1401〜1800m', 'マイラー系', '7〜9番人気', 'なし',   '127.3%', '67%（6/9年）', '2026-03-19'],
        ['B（福島ダ）★', '福島ダート1700m', 'スタミナ系', '10番人気以上', '稍・不良', '147.8%', '78%（7/9年）', '2026-03-19'],
        ['C（新潟芝）', '新潟芝2000m', 'スタミナ系', '10番人気以上', '重・4月・7月', '129.9%', '67%（6/9年）', '2026-03-19'],
        ['D（中京芝）', '中京芝1200m/1400m', '速力系', '7〜9番人気', '不良のみ', '182.8%', '60%（6/10年）', '2026-03-19'],
    ],
    col_widths=[2.0, 3.5, 2.5, 2.5, 2.5, 1.8, 2.5, 2.5]
)

add_heading(doc, '4.3 荒れ指数別馬券戦略', 2)
add_table(doc,
    ['荒れ指数帯', '戦略'],
    [
        ['〜50',   '単勝◎のみ'],
        ['50〜70', '単勝◎ ＋ 馬連◎○（○はエッジ+0.03以上・最大2頭）'],
        ['70〜85', '単勝◎のみ（馬連は見送り）'],
        ['85〜',   '全体見送りも選択肢（馬連なし）'],
    ],
    col_widths=[4.0, 12.0]
)

# ===== 5. 却下済みルール =====
doc.add_paragraph()
add_heading(doc, '5. 却下済みルール一覧', 1)
add_table(doc,
    ['ルール名', '却下理由', '検証日'],
    [
        ['○以下エッジルール',
         'ROI129.3%・黒字1/3年（2024年0/25全外れ）。主戦場1600〜2000mで47.1%赤字',
         '2026-04-25'],
        ['堅め決着ルール（前日悪→当日良）',
         '1番人気回収率56.1%（通常79.6%より-23.5%）。確認バイアス3/208件のみ',
         '2026-04-17'],
        ['条件B人気拡張（7〜9番人気追加）',
         'ROI88.0%/黒字3/11年（現行10人気以上147.8%より大幅悪化）',
         '2026-04-17'],
        ['BF軸重み2倍',
         'バックテストで-1.07%悪化確認',
         '2026-03-19'],
        ['G1初出走フラグ',
         '統計的有意差なし（バックテスト不採用）',
         '2026-04-16'],
        ['G3少キャリアフラグ',
         '統計的有意差なし（バックテスト不採用）',
         '2026-04-16'],
        ['MK軸3〜8倍加点（×1.05）',
         '2人気→1人気への誤誘導を引き起こした',
         '2026-03-17'],
        ['MK軸薄人気フラグ（×0.80）',
         '的中率-13.6%悪化',
         '2026-03-17'],
    ],
    col_widths=[4.5, 10.0, 2.5]
)

# ===== 6. 開発ロードマップ =====
doc.add_paragraph()
add_heading(doc, '6. 開発ロードマップ', 1)
add_para(doc, '最終目標：馬券種・投資額・購入判定の全自動推奨（Phase5）', bold=True)
add_table(doc,
    ['フェーズ', '内容', '期間目安', '状態'],
    [
        ['Phase1', 'エッジ値・◎○▲提示・単勝シミュレーション廃止・荒れ指数戦略', '〜2026-05', '✅ 完了'],
        ['Phase2', '馬連・ワイド期待値計算レイヤー追加（G2専用モデルと並行）',      '2026-06中', '未着手'],
        ['Phase3', 'ハーフケリー基準の投資額算出（騎手×コース相性精緻化と並行）',  '2026-07中', '未着手'],
        ['Phase4', '馬券種別バックテスト（重賞・条件戦）',                          '2026-08中', '未着手'],
        ['Phase5', '推奨ロジック統合・全自動推奨完成',                              '2026-09中', '未着手'],
    ],
    col_widths=[2.5, 8.5, 3.0, 2.0]
)
doc.add_paragraph()
add_heading(doc, '現在地（2026-05-28）', 2)
add_para(doc, '・Phase1完了。エッジ値+0.06閾値・穴馬戦略・荒れ指数別戦略が稼働中。')
add_para(doc, '・CF軸リーク排除（アプローチC）完了（5/19実装・真のROI133.1%確定）。')
add_para(doc, '・並行課題: G2専用モデル検討 / 騎手×コース相性精緻化 / 補正案Z（G2穴馬戦略優先）検証。')
add_para(doc, '・2200m超サンプル累計: 3R（目標30R超でベット判断移行）。')

# ===== 7. 運用フロー =====
doc.add_paragraph()
add_heading(doc, '7. 運用フロー', 1)

add_heading(doc, '7.1 週次フロー（土曜・日曜 9:00〜）', 2)
add_table(doc,
    ['ステップ', '時刻', '内容', '使用ツール'],
    [
        ['① 天気予報取得', '9:00〜',  '開催全会場の午後天気・降水確率取得。50%以上で馬場悪化リスク警告',   'Playwright MCP → weather_checker.py'],
        ['② 馬場状況確認', '9:30〜',  'JRA公式3ページからクッション値・馬場状態取得',                       'Playwright MCP → JRA公式'],
        ['③ 条件フィルター', '9:30〜', '条件A〜Dの該当馬を馬場状態適用後に抽出・candidates保存',            'result_checker.py --save'],
        ['④ 重賞予測',      '9:30〜', '全重賞の出馬表・オッズを取得してエッジ値・荒れ指数を算出',           'grade_race_predictor.py'],
        ['⑤ 結果照合',      'レース後', '1着馬確認・P&L計算・エッジ精度記録',                              'result_checker.py'],
    ],
    col_widths=[3.0, 1.8, 6.5, 5.0]
)

add_heading(doc, '7.2 月次メンテ（毎月末）', 2)
add_table(doc,
    ['手順', '内容', '備考'],
    [
        ['1. CSV確認',    'data/race/ 配下の新規CSVを確認。DBの最新日付より新しいファイルを対象に',  'TARGET JV出力・cp932エンコード'],
        ['2. レース結果INSERT', '重複チェック付きでrace_resultsにINSERT',                          'update_race_results.py または月次スクリプト'],
        ['3. 血統UPSERT',  'data/pedigree/ の最新血統CSVをblood_mapにUPSERT',                      '馬名をキーに上書き'],
        ['4. インデックス確認', 'idx_race_date / idx_rr_race / idx_rr_horse の存在確認',           'CREATE INDEX IF NOT EXISTS で冪等実行可'],
        ['5. VACUUM',      'SQLite VACUUM でDBサイズ圧縮',                                         '通常5〜15MB削減'],
        ['6. CLAUDE.md確認', '文字数40,000超なら作業ログをCLAUDE_archive.mdへ移動',               '確定ルール・DB情報は必ず保持'],
    ],
    col_widths=[3.5, 7.5, 5.0]
)

add_heading(doc, '7.3 データ更新手順（TARGET JV → DB）', 2)
add_table(doc,
    ['手順', '内容'],
    [
        ['1', r'TARGET JVからCSV出力 → data\race\ に配置（ファイル名: 結果yyyyMMddMMdd.csv）'],
        ['2', '血統CSV最新版を data\\pedigree\\ に配置（ファイル名: yyyyMMdd血統.csv）'],
        ['3', 'レース結果INSERTスクリプト実行（重複チェック付き）'],
        ['4', '血統UPSERTスクリプト実行'],
        ['5', 'pkl再生成: CourseMASTERv73.load_data() → build_statistics() → save()'],
        ['6', 'prediction_app.py 再起動'],
    ],
    col_widths=[1.0, 15.0]
)

add_heading(doc, '7.4 既知の制限', 2)
add_table(doc,
    ['項目', '内容'],
    [
        ['父馬名マッチ率', '75.5%（古馬・欠損でblood_map不一致 → BF=50でフォールバック）'],
        ['FR軸',           '前走馬体重データ非対応のため常に50点固定'],
        ['PD軸',           '通過順カラム欠損時は50点'],
        ['pkl',            '現行はv7.0 pkl（v7.3エンジンが後方互換で読み込み）'],
        ['SSH/VPS',        'Claude実行環境IPがConoHaセキュリティグループでブロック → 手動デプロイ'],
    ],
    col_widths=[4.0, 12.0]
)

# 保存
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f'保存完了: {OUT_PATH}')
print(f'ファイルサイズ: {os.path.getsize(OUT_PATH):,} bytes')
