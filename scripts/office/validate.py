# -*- coding: utf-8 -*-
"""Word文書の基本検証スクリプト"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

def validate(path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(path)
    print(f"検証対象: {os.path.basename(path)}")

    # 見出し抽出
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    print(f"  見出し数: {len(headings)}")
    for h in headings:
        print(f"    - {h}")

    # テーブル確認
    print(f"  テーブル数: {len(doc.tables)}")

    # 7セクション確認
    required = ['概要', 'データ現況', 'エンジン仕様', '採用済み', '却下済み', 'ロードマップ', '運用フロー']
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    missing = [r for r in required if r not in full_text]
    if missing:
        print(f"  NG: 不足セクション {missing}")
        return False
    print(f"  OK: 全7セクション確認")
    print("検証完了: PASS")
    return True

if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else None
    if not path:
        base = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        path = os.path.join(base, '仕様書', 'COURSE_MASTER_仕様書_20260528.docx')
    if not os.path.exists(path):
        print(f"ファイルが見つかりません: {path}")
        sys.exit(1)
    ok = validate(path)
    sys.exit(0 if ok else 1)
