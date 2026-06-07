# -*- coding: utf-8 -*-
"""COURSE_MASTER_仕様書_v73.docx の構造検証スクリプト"""
import sys
import os

try:
    from docx import Document
except ImportError:
    print("python-docx が未インストールです。pip install python-docx を実行してください。")
    sys.exit(1)

DOCX_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "COURSE_MASTER_仕様書_v73.docx")

def validate():
    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: ファイルが存在しません: {DOCX_PATH}")
        sys.exit(1)

    size = os.path.getsize(DOCX_PATH)
    if size == 0:
        print("ERROR: ファイルサイズが0バイトです")
        sys.exit(1)

    print(f"ファイルサイズ: {size:,} バイト ({size/1024:.1f} KB)")

    try:
        doc = Document(DOCX_PATH)
    except Exception as e:
        print(f"ERROR: Wordファイルを開けません: {e}")
        sys.exit(1)

    paragraphs = doc.paragraphs
    tables = doc.tables

    print(f"段落数: {len(paragraphs)}")
    print(f"テーブル数: {len(tables)}")

    # 章見出し確認
    headings = [p.text for p in paragraphs if p.style and p.style.name and p.style.name.startswith("Heading")]
    print(f"\n【見出し一覧】 ({len(headings)}件)")
    for h in headings:
        print(f"  {h}")

    # 期待値チェック
    errors = []
    if len(paragraphs) < 50:
        errors.append(f"段落数が少なすぎます（{len(paragraphs)} < 50）")
    if len(tables) < 10:
        errors.append(f"テーブル数が少なすぎます（{len(tables)} < 10）")
    if len(headings) < 15:
        errors.append(f"見出し数が少なすぎます（{len(headings)} < 15）")

    # 必須章の存在確認
    required = ["1. システム概要", "2. データ基盤", "3. 12軸スコアリング",
                "4. 4チーム合議制", "5. エッジ値計算", "6. バックテスト",
                "7. ADI", "8. 週次運用", "9. ファイル構成", "10. 今後の開発"]
    for req in required:
        found = any(req in h for h in headings)
        if not found:
            errors.append(f"必須章が見つかりません: {req}")

    if errors:
        print("\n【エラー】")
        for e in errors:
            print(f"  ✗ {e}")
        sys.exit(1)
    else:
        print("\n✅ 検証OK — ファイルは正常に開けます")

if __name__ == "__main__":
    validate()
