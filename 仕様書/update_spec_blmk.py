# -*- coding: utf-8 -*-
"""
update_spec_blmk.py
COURSE_MASTER_仕様書_v73.docx 更新スクリプト

追記1: 第6章末尾に 6.3（時系列リーク検証）・6.4（BL/MK検証）セクションを追加
追記2: 第10章 優先度高・中・保留テーブルに新規課題を行追加
"""
import sys
import io
import os
import copy

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, 'COURSE_MASTER_仕様書_v73.docx')
NS        = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ── ユーティリティ ────────────────────────────────────────────

def make_h2_elem(template_h2_paragraph, text):
    """
    既存の Heading 2 段落を deep copy してテキストだけ置き換えた要素を返す。
    スタイル属性がすべて保持されるため、スタイルID問題を回避できる。
    """
    new_p = copy.deepcopy(template_h2_paragraph._element)
    # w:r（run）要素を全削除してからテキスト用 run を追加
    for r in list(new_p.findall(f'{{{NS}}}r')):
        new_p.remove(r)
    r_elem = OxmlElement('w:r')
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    # 先頭・末尾の空白を保持する属性
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    new_p.append(r_elem)
    return new_p


def make_table_elem(template_tbl_element, data_rows):
    """
    既存テーブルの tblPr（枠線設定）をコピーし、data_rows のデータで
    新しいテーブル OXML 要素を構築して返す。
    data_rows[0] がヘッダー行（太字）となる。
    """
    new_tbl = OxmlElement('w:tbl')

    # tblPr コピー（枠線スタイルを維持）
    tblPr_src = template_tbl_element.find(f'{{{NS}}}tblPr')
    if tblPr_src is not None:
        new_tbl.append(copy.deepcopy(tblPr_src))

    # tblGrid（列定義）
    ncols = len(data_rows[0])
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(ncols):
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    new_tbl.append(tblGrid)

    # 行と各セルを構築
    for row_idx, row_data in enumerate(data_rows):
        tr = OxmlElement('w:tr')
        for cell_text in row_data:
            tc = OxmlElement('w:tc')
            p  = OxmlElement('w:p')

            # ヘッダー行（row_idx=0）は太字
            if row_idx == 0:
                pPr = OxmlElement('w:pPr')
                p.append(pPr)
                r_elem = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr.append(b)
                r_elem.append(rPr)
            else:
                r_elem = OxmlElement('w:r')

            t_elem = OxmlElement('w:t')
            t_elem.text = str(cell_text)
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_elem.append(t_elem)
            p.append(r_elem)
            tc.append(p)
            tr.append(tc)
        new_tbl.append(tr)

    return new_tbl


def add_table_row(docx_table, row_data):
    """既存テーブルに 1 行追加してセルテキストを設定する"""
    new_row = docx_table.add_row()
    for j, text in enumerate(row_data):
        new_row.cells[j].text = str(text)


def insert_before(ref_elem, new_elem):
    """ref_elem の直前に new_elem を挿入（lxml の addprevious を利用）"""
    ref_elem.addprevious(new_elem)


# ── ドキュメント読み込み ──────────────────────────────────────

print(f"■ ドキュメント読み込み: {DOCX_PATH}")
size_before = os.path.getsize(DOCX_PATH)
print(f"  更新前サイズ: {size_before:,} バイト ({size_before/1024:.1f} KB)")

doc  = Document(DOCX_PATH)
body = doc.element.body

# テンプレート用に既存 Heading 2 段落を取得
template_h2_para = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2':
        template_h2_para = p
        break

if template_h2_para is None:
    print("ERROR: Heading 2 段落が見つかりません")
    sys.exit(1)

# テンプレート用に既存テーブル要素を取得（枠線スタイルのコピー用）
template_tbl_elem = doc.tables[0]._tbl


# ── 挿入位置を特定（7章見出し直前） ─────────────────────────

ch7_elem = None
for child in body:
    if child.tag.split('}')[-1] != 'p':
        continue
    text = ''.join(c.text or '' for c in child.iter()
                   if hasattr(c, 'text') and c.text).strip()
    if text.startswith('7.') and 'ADI' in text:
        ch7_elem = child
        break

if ch7_elem is None:
    print("ERROR: 第7章見出しが見つかりません")
    sys.exit(1)

print(f"  挿入位置確認 OK")


# ── 追記1: 第6章末尾に 6.3・6.4 セクション追加 ────────────────

# 6.3 見出し
h63_elem = make_h2_elem(template_h2_para, '6.3 時系列リーク検証（2026-05-14実施）')
insert_before(ch7_elem, h63_elem)

# 6.3 テーブル
leak_data = [
    ['項目', '内容'],
    ['実施日', '2026-05-14'],
    ['検証内容', '統計マスターの時系列リーク定量化（leaktest_grade_races.py）'],
    ['条件A（現行）', '的中率35.85% / ROI165.5%'],
    ['条件B（リーク排除）', '的中率22.64% / ROI108.1%'],
    ['リーク量', '的中率+13.21%pt / ROI+57.4%'],
    ['主要リーク源', 'CF軸（血統CSV）が約11%pt・JT+SPD軸が約2%pt'],
    ['対応', '年次カットオフ統計4本生成（stats_cutoff_*.json）・ルーター関数実装（backtest_utils.py）'],
    ['真のベースライン', '的中率22.64% / ROI108.1%（今後のバックテスト評価基準）'],
]
insert_before(ch7_elem, make_table_elem(template_tbl_elem, leak_data))

# 6.4 見出し
h64_elem = make_h2_elem(template_h2_para, '6.4 BL/MK軸依存度検証（2026-05-15実施）')
insert_before(ch7_elem, h64_elem)

# 6.4 テーブル
blmk_data = [
    ['項目', '内容'],
    ['実施日', '2026-05-15'],
    ['検証内容', 'BL・MK軸除外時の的中率・ROI比較（4パターン・318R・時系列分離統計）'],
    ['P1現行（全軸）', '的中率22.64% / ROI108.1% / ◎平均人気3.86位'],
    ['P2（BL除外）', '的中率21.70% / ROI108.2% / ◎平均人気4.31位'],
    ['P3（MK除外）', '的中率21.38% / ROI106.3% / ◎平均人気4.31位'],
    ['P4（BL+MK除外）', '的中率19.50% / ROI106.0% / ◎平均人気4.84位'],
    ['結論', '現行維持。BL/MK除外でROI低下・的中率低下を確認。重み×0.3は妥当'],
    ['例外', 'G2限定でBL/MK除外がROI改善（62.4%→90.8%・96R・継続観察中）'],
]
insert_before(ch7_elem, make_table_elem(template_tbl_elem, blmk_data))

print("  ✓ 第6章: 6.3・6.4 セクション追加完了")


# ── 追記2: 第10章テーブルに行追加 ───────────────────────────

# Table[18] = 優先度高（課題|内容|期待効果）
add_table_row(doc.tables[18], [
    'CF軸リーク排除（アプローチC）',
    'race_resultsの年次カットオフ集計でhorse_statsをstats_cutoff_*.jsonに追加し、CF軸の完全な時系列分離を実現する',
    '現状で約11%ptのリーク排除。真の予測精度の確定',
])
print("  ✓ 第10章 優先度高: CF軸リーク排除を追加")

# Table[19] = 優先度中（課題|内容|期待効果）
add_table_row(doc.tables[19], [
    '騎手×コース相性の精緻化',
    'JT軸を騎手全体勝率から「騎手×コース×距離」の組み合わせ勝率に変更する（jockey_specialtyテーブルを活用）',
    'コース適性の予測精度向上・JT軸の情報量増加',
])
add_table_row(doc.tables[19], [
    'G2専用モデルの検討',
    'BL/MK除外でG2 ROIが62.4%→90.8%と改善（96R）。G2に特化した軸重みセットを検討',
    'G2予測精度向上（現行62.4% → 目標90%超）',
])
print("  ✓ 第10章 優先度中: 騎手×コース精緻化・G2専用モデルを追加")

# Table[20] = 保留（課題|保留理由|再評価条件）
add_table_row(doc.tables[20], [
    'maidenレースプロファイリング',
    '過学習リスク・サンプル不足（未出走馬の成績予測）',
    'サンプル200R以上蓄積後に再評価',
])
add_table_row(doc.tables[20], [
    'calibration / confidence推定',
    '予測確率の信頼区間推定。実装コスト大',
    'エンジン安定後（ROI150%以上が3年連続）に検討',
])
print("  ✓ 第10章 保留: maiden・calibrationを追加")


# ── 保存 ─────────────────────────────────────────────────────

doc.save(DOCX_PATH)
size_after = os.path.getsize(DOCX_PATH)
print(f"\n  保存完了: {DOCX_PATH}")
print(f"  更新後サイズ: {size_after:,} バイト ({size_after/1024:.1f} KB)")
print(f"  増加量: +{size_after - size_before:,} バイト")
print("■ 完了")
